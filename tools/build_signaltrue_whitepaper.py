from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "assets" / "whitepaper"
OUT_DIR = ROOT / "artifacts" / "whitepaper"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "SignalTrue_From_Work_Activity_to_Early_Signals_2026.docx"

BLUE = "3B9FF3"
DARK = "1E3A5F"
INK = "13233A"
MUTED = "60758A"
PALE = "EAF4FC"
PALE2 = "F4F8FB"
GREEN = "218C74"
GOLD = "C58B2A"
RED = "B44B4B"
WHITE = "FFFFFF"
LINE = "D7E3ED"

FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def pil_font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)


def wrapped_lines(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = word if not current else current + " " + word
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, xy, text, font, fill, max_width, line_gap=8, anchor="la"):
    x, y = xy
    lines = wrapped_lines(draw, text, font, max_width)
    h = font.size + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill, anchor=anchor)
        y += h
    return y


def make_cover_graphic(path):
    img = Image.new("RGBA", (1300, 920), (255, 255, 255, 0))
    d = ImageDraw.Draw(img)
    center = (650, 455)
    for radius, width, alpha in [(300, 6, 70), (215, 8, 105), (130, 10, 150)]:
        d.ellipse(
            (center[0] - radius, center[1] - radius, center[0] + radius, center[1] + radius),
            outline=(59, 159, 243, alpha),
            width=width,
        )
    points = [
        (425, 260, 24),
        (765, 215, 17),
        (910, 470, 26),
        (695, 675, 20),
        (420, 630, 15),
        (650, 455, 34),
    ]
    for x, y, r in points:
        d.ellipse((x-r, y-r, x+r, y+r), fill=(59, 159, 243, 230))
    for a, b in [(0, 5), (1, 5), (2, 5), (3, 5), (4, 5), (0, 1), (1, 2), (2, 3), (3, 4)]:
        d.line((points[a][0], points[a][1], points[b][0], points[b][1]), fill=(30, 58, 95, 70), width=3)
    d.rounded_rectangle((385, 735, 915, 825), 22, fill=(234, 244, 252, 235))
    d.text((650, 780), "TEAM-LEVEL SIGNALS, NOT INDIVIDUAL SCORES", font=pil_font(25, True), fill="#1E3A5F", anchor="mm")
    img.save(path)


def make_event_mix(path):
    img = Image.new("RGB", (1500, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 60), "Pilot event mix", font=pil_font(46, True), fill="#13233A")
    d.text((80, 120), "Aggregate production snapshot, June 11, 2026", font=pil_font(24), fill="#60758A")
    box = (95, 205, 595, 705)
    d.ellipse(box, fill="#EAF4FC")
    start = -90
    outlook = 847 / 1090 * 360
    d.pieslice(box, start=start, end=start + outlook, fill="#3B9FF3")
    d.pieslice(box, start=start + outlook, end=270, fill="#1E3A5F")
    d.ellipse((215, 325, 475, 585), fill="white")
    d.text((345, 420), "1,090", font=pil_font(56, True), fill="#13233A", anchor="mm")
    d.text((345, 475), "work events", font=pil_font(25), fill="#60758A", anchor="mm")
    rows = [
        ("Outlook calendar records", "847", "77.7%", "#3B9FF3"),
        ("Microsoft Teams records", "243", "22.3%", "#1E3A5F"),
        ("Approx. unique meetings", "251", "deduplicated IDs", "#218C74"),
    ]
    y = 245
    for label, value, note, color in rows:
        d.rounded_rectangle((720, y, 1405, y + 125), 20, fill="#F4F8FB", outline="#D7E3ED", width=2)
        d.ellipse((755, y + 42, 785, y + 72), fill=color)
        d.text((815, y + 32), label, font=pil_font(27, True), fill="#13233A")
        d.text((815, y + 78), note, font=pil_font(21), fill="#60758A")
        d.text((1355, y + 60), value, font=pil_font(45, True), fill=color, anchor="rm")
        y += 150
    img.save(path, quality=95)


def make_readiness(path):
    img = Image.new("RGB", (1500, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 60), "Current data readiness", font=pil_font(46, True), fill="#13233A")
    d.text((80, 120), "Daily metric snapshots are being created, but most do not yet contain usable event data.", font=pil_font(24), fill="#60758A")
    total, usable, empty = 699, 108, 591
    x0, y0, w, h = 100, 235, 1300, 105
    usable_w = int(w * usable / total)
    d.rounded_rectangle((x0, y0, x0 + w, y0 + h), 24, fill="#E7EDF2")
    d.rounded_rectangle((x0, y0, x0 + usable_w, y0 + h), 24, fill="#218C74")
    d.text((x0 + usable_w / 2, y0 + h / 2), "108 with events", font=pil_font(25, True), fill="white", anchor="mm")
    d.text((x0 + usable_w + (w - usable_w) / 2, y0 + h / 2), "591 without contributing events", font=pil_font(25, True), fill="#60758A", anchor="mm")
    cards = [
        ("699", "daily snapshots", BLUE),
        ("15%", "contain events", GREEN),
        ("0", "validated weekly strain records", RED),
        ("0", "measured intervention outcomes", GOLD),
    ]
    card_w = 300
    for i, (value, label, color) in enumerate(cards):
        x = 100 + i * 330
        d.rounded_rectangle((x, 430, x + card_w, 665), 20, fill="#F4F8FB", outline="#D7E3ED", width=2)
        d.text((x + 150, 510), value, font=pil_font(58, True), fill="#" + color, anchor="mm")
        draw_wrapped(d, (x + 150, 570), label, pil_font(22, True), "#13233A", 250, 5, anchor="ma")
    img.save(path, quality=95)


def make_pipeline(path):
    img = Image.new("RGB", (1600, 780), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 55), "How SignalTrue turns metadata into a team signal", font=pil_font(44, True), fill="#13233A")
    d.text((80, 110), "The intended method is transparent, baseline-based, and action-oriented.", font=pil_font(24), fill="#60758A")
    stages = [
        ("1", "Observe", "Meeting times, durations, participation and messaging timestamps"),
        ("2", "Aggregate", "Convert restricted event metadata into team-level daily metrics"),
        ("3", "Learn normal", "Build a 42-day team baseline using median and MAD"),
        ("4", "Detect change", "Identify sustained deviation and attach a confidence level"),
        ("5", "Act and recheck", "Recommend a reversible change and measure again after 14 days"),
    ]
    y = 225
    box_w, gap = 270, 35
    for i, (n, title, body) in enumerate(stages):
        x = 55 + i * (box_w + gap)
        d.rounded_rectangle((x, y, x + box_w, y + 395), 24, fill="#F4F8FB", outline="#D7E3ED", width=2)
        d.ellipse((x + 92, y + 30, x + 178, y + 116), fill="#3B9FF3")
        d.text((x + 135, y + 73), n, font=pil_font(34, True), fill="white", anchor="mm")
        d.text((x + 135, y + 158), title, font=pil_font(29, True), fill="#1E3A5F", anchor="mm")
        draw_wrapped(d, (x + 135, y + 210), body, pil_font(22), "#60758A", 220, 8, anchor="ma")
        if i < len(stages) - 1:
            ax = x + box_w + 8
            d.line((ax, y + 198, ax + 20, y + 198), fill="#3B9FF3", width=5)
            d.polygon([(ax + 20, y + 188), (ax + 35, y + 198), (ax + 20, y + 208)], fill="#3B9FF3")
    d.rounded_rectangle((310, 665, 1290, 735), 18, fill="#EAF4FC")
    d.text((800, 700), "Important: today, steps 1-2 are partially operational; steps 3-5 still need a complete validation dataset.", font=pil_font(24, True), fill="#1E3A5F", anchor="mm")
    img.save(path, quality=95)


def make_example(path):
    img = Image.new("RGB", (1600, 880), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 55), "Illustrative example: detecting a change from a team's own baseline", font=pil_font(42, True), fill="#13233A")
    d.text((80, 110), "Synthetic example for explanation only. It is not a customer result or validated prediction.", font=pil_font(23), fill="#B44B4B")
    left, top, right, bottom = 130, 215, 1480, 720
    for i in range(6):
        y = top + i * (bottom-top) / 5
        d.line((left, y, right, y), fill="#E1E8EE", width=2)
        d.text((95, y), str(100 - i*20), font=pil_font(20), fill="#60758A", anchor="rm")
    labels = ["W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8", "W9", "W10"]
    for i, label in enumerate(labels):
        x = left + i * (right-left)/(len(labels)-1)
        d.text((x, bottom+35), label, font=pil_font(20), fill="#60758A", anchor="ma")
    baseline = [49, 51, 48, 52, 50, 49, 51, 50, 49, 50]
    current = [50, 48, 51, 49, 52, 60, 69, 77, 82, 78]
    def xy(vals):
        pts=[]
        for i,v in enumerate(vals):
            x=left+i*(right-left)/(len(vals)-1)
            y=bottom-(v/100)*(bottom-top)
            pts.append((x,y))
        return pts
    bpts, cpts = xy(baseline), xy(current)
    d.line(bpts, fill="#AEBCC8", width=5)
    d.line(cpts, fill="#3B9FF3", width=8)
    for x,y in cpts: d.ellipse((x-7,y-7,x+7,y+7), fill="#3B9FF3")
    threshold_y = bottom - .70*(bottom-top)
    d.line((left, threshold_y, right, threshold_y), fill="#C58B2A", width=3)
    d.text((right-10, threshold_y-16), "illustrative alert threshold", font=pil_font(20, True), fill="#C58B2A", anchor="ra")
    d.rounded_rectangle((760, 250, 1440, 345), 18, fill="#EAF4FC")
    d.text((790, 275), "Example interpretation", font=pil_font(24, True), fill="#1E3A5F")
    d.text((790, 313), "The pattern has moved beyond normal weekly variation for three consecutive weeks.", font=pil_font(21), fill="#60758A")
    d.text((130, 800), "Signal index", font=pil_font(20), fill="#60758A")
    d.line((1130, 800, 1190, 800), fill="#3B9FF3", width=7)
    d.text((1205, 800), "observed pattern", font=pil_font(20), fill="#13233A", anchor="lm")
    d.line((1300, 800, 1360, 800), fill="#AEBCC8", width=5)
    d.text((1375, 800), "baseline", font=pil_font(20), fill="#13233A", anchor="lm")
    img.save(path, quality=95)


def make_privacy(path):
    img = Image.new("RGB", (1600, 820), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 55), "Privacy architecture: useful visibility without individual scoring", font=pil_font(42, True), fill="#13233A")
    columns = [
        (80, "Restricted event layer", ["Timestamps", "Meeting duration", "Participation links", "Source and event type"], BLUE),
        (585, "Aggregation controls", ["Minimum 8 active people", "Minimum 5 contributors per metric", "Concentration check", "Missing data stays missing"], GREEN),
        (1090, "Customer-visible output", ["Team trends", "Metric drivers", "Confidence", "No individual ranking"], DARK),
    ]
    for x, title, items, color in columns:
        d.rounded_rectangle((x, 185, x+430, 680), 26, fill="#F4F8FB", outline="#D7E3ED", width=2)
        d.rounded_rectangle((x, 185, x+430, 280), 26, fill="#"+color)
        d.rectangle((x, 235, x+430, 280), fill="#"+color)
        d.text((x+215, 233), title, font=pil_font(27, True), fill="white", anchor="mm")
        yy=340
        for item in items:
            d.ellipse((x+45, yy-8, x+65, yy+12), fill="#"+color)
            d.text((x+90, yy+2), item, font=pil_font(23), fill="#13233A", anchor="lm")
            yy += 75
    for x in (535, 1040):
        d.line((x, 430, x+30, 430), fill="#3B9FF3", width=6)
        d.polygon([(x+30, 418),(x+50,430),(x+30,442)], fill="#3B9FF3")
    d.rounded_rectangle((255, 720, 1345, 785), 18, fill="#FFF4E0")
    d.text((800, 752), "Before publication, stored identifiers and retention behavior must be aligned with this promise.", font=pil_font(24, True), fill="#7C5B1F", anchor="mm")
    img.save(path, quality=95)


def make_roadmap(path):
    img = Image.new("RGB", (1600, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 55), "Evidence roadmap", font=pil_font(44, True), fill="#13233A")
    d.text((80, 110), "The publication today starts a research program; it does not pretend the program is finished.", font=pil_font(24), fill="#60758A")
    phases = [
        ("NOW", "Methods paper", "Explain the method, pilot data, privacy design and limitations", BLUE),
        ("0-3 months", "Founding cohort", "Clean data, collect validated surveys and execution outcomes", GREEN),
        ("4-9 months", "Lead-lag study", "Test which signals move before independent outcomes", GOLD),
        ("9-15 months", "Intervention evidence", "Measure which reversible actions improve target metrics", DARK),
    ]
    y=235
    d.line((200, y+75, 1400, y+75), fill="#D7E3ED", width=10)
    for i,(period,title,body,color) in enumerate(phases):
        x=200+i*400
        d.ellipse((x-28,y+47,x+28,y+103), fill="#"+color)
        d.text((x,y), period, font=pil_font(23, True), fill="#"+color, anchor="ma")
        d.rounded_rectangle((x-165,y+150,x+165,y+470),22,fill="#F4F8FB",outline="#D7E3ED",width=2)
        d.text((x,y+205),title,font=pil_font(27,True),fill="#1E3A5F",anchor="ma")
        draw_wrapped(d,(x,y+280),body,pil_font(21),"#60758A",270,8,anchor="ma")
    img.save(path, quality=95)


def generate_graphics():
    files = {
        "cover": ASSET_DIR / "cover_signal.png",
        "event_mix": ASSET_DIR / "pilot_event_mix.png",
        "readiness": ASSET_DIR / "data_readiness.png",
        "pipeline": ASSET_DIR / "signal_pipeline.png",
        "example": ASSET_DIR / "illustrative_baseline.png",
        "privacy": ASSET_DIR / "privacy_architecture.png",
        "roadmap": ASSET_DIR / "evidence_roadmap.png",
    }
    make_cover_graphic(files["cover"])
    make_event_mix(files["event_mix"])
    make_readiness(files["readiness"])
    make_pipeline(files["pipeline"])
    make_example(files["example"])
    make_privacy(files["privacy"])
    make_roadmap(files["roadmap"])
    return files


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_picture_with_alt(run, path, width, alt_text):
    picture = run.add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)
    return picture


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, widths_inches):
    table.autofit = False
    total_dxa = sum(round(w * 1440) for w in widths_inches)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            w = round(widths_inches[idx] * 1440)
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(w))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_paragraph_border(paragraph, color=BLUE, size=16, space=6, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def shade_paragraph(paragraph, fill=PALE2):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_hyperlink(paragraph, text, url, color=BLUE, underline=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    r_pr.append(sz)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in [
        ("Title", 30, DARK, 0, 8),
        ("Subtitle", 15, MUTED, 0, 12),
        ("Heading 1", 16, DARK, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SIGNALTRUE  /  METHODS & PILOT DATA WHITEPAPER")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(p, color=LINE, size=8, space=3, side="bottom")
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_cover(doc, graphics):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    logo = ROOT / "public" / "logo512.png"
    if logo.exists() and logo.stat().st_size > 0:
        add_picture_with_alt(p.add_run(), logo, 0.62, "SignalTrue logo")
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_after = Pt(38)
    set_run_font(brand.add_run("SignalTrue"), size=16, color=DARK, bold=True)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(15)
    set_run_font(kicker.add_run("METHODS & PILOT DATA WHITEPAPER"), size=10, color=BLUE, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("From Work Activity\nto Early Signals")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("How SignalTrue turns collaboration metadata into team-level indicators without reading employee content")
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_before = Pt(6)
    pic.paragraph_format.space_after = Pt(18)
    add_picture_with_alt(
        pic.add_run(),
        graphics["cover"],
        4.55,
        "Concentric team-signal illustration showing observations becoming aggregated indicators.",
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(2)
    meta.paragraph_format.space_after = Pt(2)
    set_run_font(meta.add_run("June 2026  |  Version 1.0"), size=10.5, color=MUTED, bold=True)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(note.add_run("Descriptive pilot evidence. Not an industry benchmark or validated prediction model."), size=9.5, color=RED, italic=True)
    first_footer = doc.sections[0].first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(first_footer.add_run("signaltrue.ai  |  Team-level work signals, metadata only"), size=8.5, color=MUTED)
    doc.add_page_break()


def add_callout(doc, title, body, color=BLUE, fill=PALE2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, fill)
    set_paragraph_border(p, color=color, size=22, space=7, side="left")
    set_run_font(p.add_run(title + "\n"), size=11, color=color, bold=True)
    set_run_font(p.add_run(body), size=11, color=INK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(text), size=8.5, color=MUTED, italic=True)


def add_figure(doc, path, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    add_picture_with_alt(p.add_run(), path, width, caption)
    add_caption(doc, caption)


def add_metric_table(doc):
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_fixed_width(table, [1.62, 1.62, 1.62, 1.62])
    set_repeat_table_header(table.rows[0])
    values = [("1,090", "work events"), ("699", "daily snapshots"), ("10", "teams configured"), ("0", "validated outcomes")]
    for i, (value, label) in enumerate(values):
        top, bottom = table.cell(0, i), table.cell(1, i)
        shade_cell(top, PALE)
        shade_cell(bottom, WHITE)
        top.vertical_alignment = bottom.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        top.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bottom.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(top.paragraphs[0].add_run(value), size=22, color=BLUE if i < 3 else RED, bold=True)
        set_run_font(bottom.paragraphs[0].add_run(label), size=9.5, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_two_column_table(doc, rows, widths=(2.0, 4.45), header=None):
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_width(table, list(widths))
    offset = 0
    if header:
        offset = 1
        for i, text in enumerate(header):
            cell = table.cell(0, i)
            shade_cell(cell, DARK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(text), size=10, color=WHITE, bold=True)
        set_repeat_table_header(table.rows[0])
    else:
        # Key-value tables use the first row as their accessible orientation row.
        set_repeat_table_header(table.rows[0])
    for r_idx, (left, right) in enumerate(rows, start=offset):
        c1, c2 = table.cell(r_idx, 0), table.cell(r_idx, 1)
        if (r_idx - offset) % 2 == 0:
            shade_cell(c1, PALE2)
            shade_cell(c2, PALE2)
        c1.vertical_alignment = c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_run_font(c1.paragraphs[0].add_run(left), size=10, color=DARK, bold=True)
        set_run_font(c2.paragraphs[0].add_run(right), size=10, color=INK)
    return table


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=11, color=INK, bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=11, color=INK)
    else:
        set_run_font(p.add_run(text), size=11, color=INK)
    return p


def add_source(doc, number, title, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run(f"{number}. {title}. "), size=9.5, color=INK)
    add_hyperlink(p, url, url)


def build_document(graphics):
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        add_running_furniture(section)
    core = doc.core_properties
    core.title = "From Work Activity to Early Signals"
    core.subject = "SignalTrue methods and pilot data whitepaper"
    core.author = "SignalTrue"
    core.keywords = "work signals, organizational analytics, privacy, metadata, team health"

    add_cover(doc, graphics)

    doc.add_heading("Executive summary", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("Organizations usually discover work-system problems late: "), size=12, color=DARK, bold=True)
    set_run_font(p.add_run("after a survey falls, delivery slips, or someone resigns. SignalTrue is designed to make earlier structural changes visible through team-level metadata such as meeting load, focus fragmentation, after-hours activity, response pressure, and collaboration patterns."), size=12, color=INK)
    add_callout(
        doc,
        "THE CENTRAL IDEA",
        "A team's normal pattern is more useful than a universal productivity score. SignalTrue learns what is normal for each team, then looks for sustained changes that leaders can investigate and address.",
    )
    doc.add_paragraph("This paper has three purposes:")
    add_bullet(doc, "Explain the SignalTrue method in plain language.")
    add_bullet(doc, "Show what the current aggregate pilot data contains and what it does not yet prove.")
    add_bullet(doc, "Invite organizations to participate in a prospective validation study.")
    add_metric_table(doc)
    add_callout(
        doc,
        "IMPORTANT LIMITATION",
        "The current dataset is suitable for demonstrating the method and identifying data-quality priorities. It is not large or complete enough to claim that SignalTrue predicts burnout, resignation, or delivery failure.",
        color=RED,
        fill="FCEEEE",
    )

    doc.add_heading("1. The visibility gap", level=1)
    doc.add_paragraph("Most organizations rely on two kinds of information:")
    add_bullet(doc, "Lagging human signals, such as engagement surveys, complaints, sickness absence, and resignations.")
    add_bullet(doc, "Lagging operational signals, such as missed deadlines, slow decisions, rework, and quality problems.")
    doc.add_paragraph("Between those two is a largely invisible layer: the operating conditions of work. Calendars become fragmented. Coordination consumes more attendee-hours. Work moves into evenings. Response demand rises. Collaboration narrows. Individually, each change may be harmless. Sustained combinations can indicate that a team is compensating for a structural problem.")
    add_callout(doc, "SIGNALTRUE'S QUESTION", "Can these operating changes be measured early enough for leaders to make a small, reversible change before the situation becomes expensive or painful?")
    doc.add_heading("What makes the question useful", level=2)
    add_two_column_table(doc, [
        ("Earlier", "It focuses on work patterns that can change before survey or business outcomes."),
        ("Structural", "It asks how work is organized, not whether individuals are working hard enough."),
        ("Explainable", "Every signal should show which metrics changed and how confident the system is."),
        ("Actionable", "The intended endpoint is a reversible team intervention, followed by measurement."),
    ])

    doc.add_heading("2. What SignalTrue measures", level=1)
    add_figure(doc, graphics["pipeline"], 6.55, "Figure 1. Intended SignalTrue measurement and action loop.")
    doc.add_paragraph("SignalTrue is designed around three families of operating conditions:")
    add_two_column_table(doc, [
        ("Demand", "Meeting load, attendee-hours, message volume, response demand, workload volatility."),
        ("Recovery", "After-hours activity, back-to-back meetings, short recovery gaps, fragmented days."),
        ("Flow", "Focus availability, collaboration breadth, reciprocity, task aging, cycle time, rework."),
    ], header=("Signal family", "Examples"))
    doc.add_heading("The baseline method", level=2)
    doc.add_paragraph("The newest SignalTrue method uses a 42-day rolling baseline. Instead of relying on the average alone, it uses the median and median absolute deviation (MAD). This makes the baseline less sensitive to unusual weeks, launches, off-sites, or one exceptionally busy day.")
    add_callout(doc, "PLAIN-LANGUAGE INTERPRETATION", "A signal is not 'this team has too many meetings.' It is closer to 'this team's coordination pattern has moved materially beyond its own normal range, and the change has persisted.'", color=GREEN, fill="EDF8F4")

    doc.add_heading("3. What SignalTrue does not measure", level=1)
    doc.add_paragraph("The intended product boundary is as important as the metrics. SignalTrue is not designed to evaluate individual productivity or infer personal psychological states.")
    add_two_column_table(doc, [
        ("Message or email content", "Not needed for the proposed team-level work-pattern method."),
        ("Keystrokes, screens or websites", "No device surveillance or application tracking."),
        ("Individual productivity", "No employee rankings, personal activity scores, or performance labels."),
        ("Clinical diagnosis", "No claim that metadata can diagnose burnout, depression, anxiety, or another condition."),
    ], header=("Excluded", "Product boundary"))

    doc.add_heading("4. Current pilot data", level=1)
    doc.add_paragraph("The following is an aggregate snapshot of the connected SignalTrue database on June 11, 2026. No company names, team names, or individual identities are included.")
    add_figure(doc, graphics["event_mix"], 6.45, "Figure 2. Current production event mix. Calendar events are stored at attendee level; approximately 251 user-independent meeting IDs were identified.")
    add_figure(doc, graphics["readiness"], 6.45, "Figure 3. Data readiness across 699 daily integration metric snapshots.")
    doc.add_heading("What is present", level=2)
    add_bullet(doc, "1,090 normalized work-event records from Microsoft Outlook and Teams.")
    add_bullet(doc, "847 calendar records, representing approximately 251 unique meeting IDs after removing attendee-specific suffixes.")
    add_bullet(doc, "243 Teams message records and 699 daily metric snapshots.")
    add_bullet(doc, "Calendar duration, meeting counts, back-to-back blocks, average gaps, after-hours messaging, and basic source coverage.")
    doc.add_heading("What is missing", level=2)
    add_bullet(doc, "No completed records from the new Engagement Strain daily, baseline, or weekly scoring pipeline.")
    add_bullet(doc, "No measured actions or intervention outcomes.")
    add_bullet(doc, "No validated survey, delivery, absence, or turnover labels linked to team-week signals.")
    add_bullet(doc, "Incomplete Teams attribution and missing thread/channel enrichment needed for collaboration metrics.")

    doc.add_heading("5. What the data can say today", level=1)
    add_two_column_table(doc, [
        ("Can support", "A transparent method demonstration, private pilot diagnostic, and design-partner research protocol."),
        ("Can support", "Descriptive observations about data coverage, meeting duration, event mix, and measurement limitations."),
        ("Cannot support", "An industry benchmark, because the current usable data is concentrated in one connected environment."),
        ("Cannot support", "A claim that a SignalTrue score predicts burnout, resignation, or delivery failure."),
        ("Cannot support", "A quantified ROI or intervention-effect claim, because no completed interventions are recorded."),
    ], header=("Evidence status", "Responsible interpretation"))
    add_callout(doc, "WHY THIS HONESTY MATTERS", "A credible whitepaper should make the boundary between observed data, illustrative examples, and future hypotheses unmistakable. Trust is part of the product.", color=GREEN, fill="EDF8F4")
    doc.add_heading("Data-quality findings", level=2)
    add_bullet(doc, "Calendar records are stored once per internal attendee. Downstream calculations must use the correct grain to avoid counting one meeting several times.")
    add_bullet(doc, "Six meeting copies exceed 12 hours; all-day and malformed events need explicit filtering.")
    add_bullet(doc, "Only 108 of 699 daily snapshots contain processed events. Missing data should remain missing rather than becoming a neutral score.")
    add_bullet(doc, "Current signal records have confidence values of zero, so they should not be presented as validated findings.")
    add_bullet(doc, "The privacy suppression system has recorded 376 events, mostly because team-size metadata is missing and reported as zero.")

    doc.add_heading("6. Example: from baseline to signal", level=1)
    add_figure(doc, graphics["example"], 6.55, "Figure 4. Synthetic explanation of a sustained deviation from a team's own baseline. Not customer evidence.")
    doc.add_paragraph("In this illustrative example, the observed pattern remains close to baseline for five weeks, then rises for three consecutive weeks. A responsible signal would include:")
    add_bullet(doc, "The metric that changed and its baseline range.")
    add_bullet(doc, "How long the deviation persisted.")
    add_bullet(doc, "The number of active contributors and integrations supporting the result.")
    add_bullet(doc, "Relevant context, such as a launch, holiday, reorganization, or incident.")
    add_bullet(doc, "A reversible action and a date for re-measurement.")

    doc.add_heading("7. Privacy without blindness", level=1)
    add_figure(doc, graphics["privacy"], 6.55, "Figure 5. Intended separation between restricted event processing and customer-visible team signals.")
    doc.add_paragraph("A privacy-preserving organizational analytics product still needs restricted event-level processing to attribute activity to the correct team and calculate distributions. The defensible promise is not that no identifiers ever enter the system. It is that identifiers are minimized and protected, individual analytics are not exposed, and only sufficiently large team aggregates become visible.")
    doc.add_heading("Publication standard", level=2)
    add_bullet(doc, "At least eight active people in a reported team-week.")
    add_bullet(doc, "At least five contributors for each displayed metric.")
    add_bullet(doc, "Suppression or a concentration warning if one person accounts for more than 40% of a sensitive metric.")
    add_bullet(doc, "Clear retention, deletion, role-based access, and benchmark opt-in rules.")
    add_bullet(doc, "No customer-visible individual activity records or individual scores.")
    add_callout(doc, "IMPLEMENTATION NOTE", "Before this privacy architecture is marketed as fully implemented, SignalTrue should remove unnecessary email fields, consistently tokenize identities, correct raw-event retention deletion, and align all public minimum-team-size statements.", color=GOLD, fill="FFF5E6")

    doc.add_heading("8. The research question that matters", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(16)
    shade_paragraph(p, PALE)
    set_run_font(p.add_run("Which changes in team work patterns provide a reliable early warning, how early do they appear, and which low-risk interventions reverse them?"), size=16, color=DARK, bold=True)
    doc.add_paragraph("This question is more valuable than another report saying that knowledge workers have too many meetings. Large vendors already publish extensive descriptive benchmarks. SignalTrue can contribute something smaller but more useful: independent evidence linking an explainable team signal to a later outcome and then to a measured intervention.")

    doc.add_heading("9. Proposed founding study", level=1)
    add_two_column_table(doc, [
        ("Cohort", "10-20 organizations and 50-150 teams."),
        ("Observation period", "16-26 weeks, long enough to build past-only baselines and test lead time."),
        ("Primary unit", "Team-week; daily events are used only to construct quality-controlled metrics."),
        ("Independent outcomes", "Validated short surveys plus delivery measures from Jira, Asana, Linear, or operational systems."),
        ("Context", "Holidays, launches, incidents, planning weeks, reorganizations, and staffing changes."),
        ("Analysis", "Within-team normalization, mixed-effects models, lag tests, holdout organizations, calibration and false-alert rates."),
    ], header=("Study element", "Proposed design"))
    doc.add_heading("Primary hypotheses", level=2)
    p = doc.add_paragraph(style="List Number")
    set_run_font(p.add_run("Recovery and focus signals will move before independent team-reported strain more reliably than raw activity volume alone."), size=11, color=INK)
    p = doc.add_paragraph(style="List Number")
    set_run_font(p.add_run("Combinations of coordination burden and execution-flow measures will provide more useful warning than a single universal score."), size=11, color=INK)
    p = doc.add_paragraph(style="List Number")
    set_run_font(p.add_run("Targeted, reversible interventions will improve the metric they are designed to change when the intervention is actually adopted."), size=11, color=INK)
    doc.add_heading("Measures", level=2)
    add_two_column_table(doc, [
        ("Employee-reported", "UWES-3 engagement and a short validated strain, exhaustion, or detachment measure."),
        ("Execution", "Cycle time, aging work, reopen rate, missed commitment, or SLA breach."),
        ("Manager validation", "Weekly signal accuracy and context review."),
        ("Intervention", "Target metric at baseline, 14 days, and 28 days, plus an adherence check."),
    ])

    doc.add_heading("10. Publication roadmap", level=1)
    add_figure(doc, graphics["roadmap"], 6.55, "Figure 6. Recommended progression from methods transparency to validated evidence.")
    doc.add_paragraph("The first flagship evidence report should be titled:")
    add_callout(doc, "BEFORE THE SURVEY MOVES", "Which team work signals changed first, how many weeks of warning they provided, and what false-alert rate leaders should expect.", color=BLUE, fill=PALE)
    doc.add_paragraph("The second should focus on action outcomes:")
    add_callout(doc, "WHAT ACTUALLY RESTORES FOCUS", "Evidence from 14- and 28-day team interventions, including what worked, what did not, and under which operating conditions.", color=GREEN, fill="EDF8F4")

    doc.add_heading("11. What leaders can do now", level=1)
    doc.add_paragraph("Even before predictive validation is complete, organizations can use the method responsibly as a structured conversation aid:")
    add_bullet(doc, "Review changes against each team's own history instead of ranking teams against one another.")
    add_bullet(doc, "Ask what changed in the work system before asking what is wrong with the people.")
    add_bullet(doc, "Choose one low-risk intervention, define the expected metric movement, and set a recheck date.")
    add_bullet(doc, "Treat a signal as a prompt for investigation, not as proof of burnout, disengagement, or poor performance.")
    add_bullet(doc, "Share the data method and privacy controls with employees and representatives before rollout.")
    doc.add_heading("Invitation: Founding Work Signals Study", level=1)
    doc.add_paragraph("SignalTrue is seeking design partners for a 16-26 week prospective study. Participating organizations will help define a privacy-respecting evidence standard for early organizational signals.")
    add_two_column_table(doc, [
        ("Participants receive", "A private data-quality audit and team baseline report."),
        ("Participants receive", "Quarterly cohort benchmarks where privacy thresholds permit."),
        ("Participants receive", "A co-designed intervention and 14-/28-day recheck."),
        ("Participants contribute", "Metadata integrations, team mapping, short pulse measures, operational outcomes, and context events."),
    ])
    cta = doc.add_paragraph()
    cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cta.paragraph_format.space_before = Pt(16)
    cta.paragraph_format.space_after = Pt(12)
    shade_paragraph(cta, DARK)
    set_run_font(cta.add_run("JOIN THE FOUNDING STUDY\n"), size=14, color=WHITE, bold=True)
    set_run_font(cta.add_run("Request a research and privacy briefing at signaltrue.ai"), size=11, color=WHITE)

    doc.add_page_break()
    doc.add_heading("Method notes and limitations", level=1)
    add_bullet(doc, "Snapshot date: June 11, 2026. Counts may change as integrations continue to sync.")
    add_bullet(doc, "Approximate unique meetings were estimated by removing attendee-specific suffixes from stored Outlook external IDs.")
    add_bullet(doc, "The pilot dataset is concentrated in Microsoft Outlook and Teams and is not representative of the broader market.")
    add_bullet(doc, "Current composite scores were not used as evidence because missing inputs, proxy reuse, incomplete enrichment, and neutral defaults require correction.")
    add_bullet(doc, "The illustrative baseline chart is synthetic and is included only to explain the method.")
    add_bullet(doc, "No individual or customer-identifying data is included in this document.")
    doc.add_heading("References", level=1)
    add_source(doc, 1, "Microsoft Work Trend Index, Breaking down the infinite workday", "https://www.microsoft.com/en-us/worklab/work-trend-index/breaking-down-infinite-workday")
    add_source(doc, 2, "Microsoft Viva Insights privacy guide", "https://learn.microsoft.com/en-us/viva/insights/advanced/privacy/privacy")
    add_source(doc, 3, "Worklytics, 2025 productivity benchmarks", "https://www.worklytics.co/resources/2025-productivity-benchmarks-knowledge-workers-teams-above-below-line")
    add_source(doc, 4, "Schaufeli et al., UWES-3 validation", "https://pmc.ncbi.nlm.nih.gov/articles/PMC6161491/")
    add_source(doc, 5, "Bakker, Demerouti and Sanz-Vergel, Job Demands-Resources theory", "https://www.annualreviews.org/content/journals/10.1146/annurev-orgpsych-120920-053933")
    add_source(doc, 6, "Microsoft Research, research-backed practices for better meetings", "https://www.microsoft.com/en-us/research/articles/research-backed-practices-for-better-meetings/")
    add_source(doc, 7, "European Data Protection Board, Guidelines 1/2026 on scientific research", "https://www.edpb.europa.eu/our-work-tools/documents/public-consultations/2026/guidelines-12026-processing-personal-data_en")
    doc.add_heading("About SignalTrue", level=1)
    doc.add_paragraph("SignalTrue is developing a privacy-conscious work-signal intelligence platform for team-level organizational visibility. The product is designed to help leaders see changes in coordination, recovery, focus, and collaboration early enough to investigate and improve how work is structured.")
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.paragraph_format.space_before = Pt(22)
    set_run_font(final.add_run("SignalTrue makes the operating conditions of work visible early enough to change them."), size=14, color=DARK, bold=True)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    graphics = generate_graphics()
    path = build_document(graphics)
    print(path)
